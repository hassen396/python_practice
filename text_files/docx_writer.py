from docx import Document
doc = Document()
doc.add_heading('My doc Title')
doc.add_paragraph('This is a real docx file created using python')
doc.save('./written.docx')